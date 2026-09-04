"""Build the NFA Word document from an NFAData object.

This module knows nothing about Ariba. It takes plain data and produces a
.docx, which is what lets the whole document pipeline be exercised offline.
"""

from __future__ import annotations

import re
from datetime import date
from decimal import Decimal
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from ..model import NFAData, VendorQuote
from ..mapping.formatting import PLACEHOLDER, date_str, inr_words, plain_amount, variance_str
from . import styles
from .docx_helpers import (
    add_page_x_of_y,
    prevent_row_split,
    remove_table_borders,
    repeat_header_row,
    set_cell_margins,
    set_col_widths,
    set_page_borders,
    shade_cell,
)

TITLE = "AWARD NOTE / NOTE FOR APPROVAL"

# Column widths in cm; each set sums to styles.CONTENT_WIDTH_CM (15.92).
GRID_WIDTHS = [1.0, 6.4, 8.5]
COMPARISON_WIDTHS = [1.3, 4.9, 2.7, 2.3, 2.7, 2.0]
#: #, Line Item, Qty, UoM - the vendor columns share what is left.
ITEM_FIXED_WIDTHS = [0.8, 4.6, 1.5, 1.1]
#: Item ID, Item Name, UOM, Qty, Last PO Price, L1 Price, Total Price, L1 Vendor.
DETAIL_WIDTHS = [3.0, 3.0, 0.9, 1.2, 1.9, 1.7, 1.8, 2.42]


def _vendor_column_widths(count: int) -> list[float]:
    remaining = styles.CONTENT_WIDTH_CM - sum(ITEM_FIXED_WIDTHS)
    return [remaining / max(count, 1)] * count


class NFABuilder:
    """Renders one NFAData into a Word document."""

    def __init__(self, data: NFAData):
        self.data = data
        self.document = Document()

    def build(self) -> Document:
        styles.apply_styles(self.document)
        section = self.document.sections[0]
        styles.setup_page(section)
        set_page_borders(section)

        self._letterhead()
        self._date()
        self._title()
        self._subject()
        self._grid()
        self._comparison()
        self._item_wise_prices()
        self._line_item_detail()
        self._justification()
        self._footer()
        self._core_properties()
        return self.document

    def save(self, path: Path) -> Path:
        path = Path(path)
        path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(path))
        return path

    # ------------------------------------------------------------------ #
    # Sections
    # ------------------------------------------------------------------ #

    def _letterhead(self) -> None:
        """Logo beside the entity name, or the name alone when no logo is set.

        An entity with no logo configured is a deliberate choice, not a missing
        field, so it gets a clean centred name rather than a red placeholder.
        A logo that *is* configured but whose file is absent is a genuine
        misconfiguration and is flagged.
        """
        entity = self.data.entity
        wants_logo = bool(entity.logo_path)

        if not wants_logo:
            table = self.document.add_table(rows=1, cols=1)
            remove_table_borders(table)
            set_col_widths(table, [styles.CONTENT_WIDTH_CM])
            self._entity_text(table.rows[0].cells[0], WD_ALIGN_PARAGRAPH.CENTER)
            return

        if not entity.show_name:
            # A wordmark logo carries the name itself, so it sits centred alone.
            table = self.document.add_table(rows=1, cols=1)
            remove_table_borders(table)
            set_col_widths(table, [styles.CONTENT_WIDTH_CM])
            cell = table.rows[0].cells[0]
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._place_logo(para, entity)
            if entity.address:
                addr = cell.add_paragraph()
                addr.alignment = WD_ALIGN_PARAGRAPH.CENTER
                addr.add_run(entity.address).font.size = Pt(9)
            return

        table = self.document.add_table(rows=1, cols=2)
        remove_table_borders(table)
        set_col_widths(table, [5.0, styles.CONTENT_WIDTH_CM - 5.0])

        logo_cell, text_cell = table.rows[0].cells
        logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        self._place_logo(logo_cell.paragraphs[0], entity)
        self._entity_text(text_cell, WD_ALIGN_PARAGRAPH.RIGHT)

    def _place_logo(self, paragraph, entity) -> None:
        if Path(entity.logo_path).exists():
            paragraph.add_run().add_picture(
                str(entity.logo_path), width=Cm(entity.logo_width_cm)
            )
        else:
            self._missing_run(paragraph, "logo file not found")

    def _entity_text(self, cell, alignment) -> None:
        entity = self.data.entity
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        name_para = cell.paragraphs[0]
        name_para.alignment = alignment
        name_run = name_para.add_run(entity.name)
        name_run.bold = True
        name_run.font.size = Pt(15) if alignment == WD_ALIGN_PARAGRAPH.CENTER else Pt(13)

        if entity.address:
            addr_para = cell.add_paragraph()
            addr_para.alignment = alignment
            addr_run = addr_para.add_run(entity.address)
            addr_run.font.size = Pt(9)

    def _date(self) -> None:
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        value = self.data.doc_date or date.today()
        para.add_run(f"Date: {date_str(value)}").bold = True

    def _title(self) -> None:
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(TITLE)
        run.bold = True
        run.underline = True
        run.font.size = Pt(13)

    def _subject(self) -> None:
        para = self.document.add_paragraph()
        para.paragraph_format.space_after = Pt(8)
        para.add_run("Subject: ").bold = True
        if self.data.subject:
            para.add_run(self.data.subject)
        else:
            self._missing_run(para, "subject not provided")

    def _grid(self) -> None:
        rows = self.data.grid
        table = self.document.add_table(rows=len(rows), cols=3)
        table.style = "Table Grid"
        set_col_widths(table, GRID_WIDTHS)
        set_cell_margins(table)

        for index, row_data in enumerate(rows):
            cells = table.rows[index].cells
            prevent_row_split(table.rows[index])

            letter_para = cells[0].paragraphs[0]
            letter_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            letter_para.add_run(row_data.letter).bold = True
            shade_cell(cells[0], styles.HEADER_FILL)

            label_run = cells[1].paragraphs[0].add_run(row_data.label)
            label_run.font.size = styles.TABLE_SIZE

            value_para = cells[2].paragraphs[0]
            if row_data.missing:
                self._missing_run(value_para, f"{row_data.label} — not provided")
            else:
                self._multiline(value_para, row_data.value, size=styles.TABLE_SIZE)

    def _comparison(self) -> None:
        self._heading("Comparison of Offers (Top 3 Participating Vendors)")

        vendors = self.data.vendors
        if not vendors:
            para = self.document.add_paragraph()
            self._missing_run(para, "No priced vendor offers available for comparison")
            return

        headers = [
            "Rank", "Vendor",
            "Basic Value (INR)", f"GST @{self._gst_rate_text()}%",
            "Total (INR)", "Variance",
        ]
        table = self.document.add_table(rows=len(vendors) + 1, cols=len(headers))
        table.style = "Table Grid"
        set_col_widths(table, COMPARISON_WIDTHS)
        set_cell_margins(table)

        header_row = table.rows[0]
        repeat_header_row(header_row)
        for index, text in enumerate(headers):
            cell = header_row.cells[index]
            shade_cell(cell, styles.HEADER_FILL)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            run.bold = True
            run.font.size = styles.TABLE_SIZE

        for offset, vendor in enumerate(vendors, start=1):
            self._comparison_row(table.rows[offset], vendor)

        self._comparison_notes()
        self._award_summary()

    def _comparison_row(self, row, vendor: VendorQuote) -> None:
        prevent_row_split(row)
        name = vendor.name + (" *" if vendor.partial_bid else "")
        values = [
            (vendor.rank_label, WD_ALIGN_PARAGRAPH.CENTER, True),
            (name, WD_ALIGN_PARAGRAPH.LEFT, vendor.rank == 1),
            (plain_amount(vendor.basic), WD_ALIGN_PARAGRAPH.RIGHT, False),
            (plain_amount(vendor.gst), WD_ALIGN_PARAGRAPH.RIGHT, False),
            (plain_amount(vendor.total), WD_ALIGN_PARAGRAPH.RIGHT, vendor.rank == 1),
            (variance_str(vendor.variance_pct), WD_ALIGN_PARAGRAPH.CENTER, False),
        ]
        for index, (text, alignment, bold) in enumerate(values):
            para = row.cells[index].paragraphs[0]
            para.alignment = alignment
            run = para.add_run(text)
            run.bold = bold
            run.font.size = styles.TABLE_SIZE

    def _comparison_notes(self) -> None:
        notes = list(self.data.report.notes)
        if any(v.partial_bid for v in self.data.vendors):
            notes.insert(0, "* Partial bid - vendor did not quote every line item.")
        for note in notes:
            para = self.document.add_paragraph()
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(note)
            run.font.size = Pt(9)
            run.italic = True

    def _award_summary(self) -> None:
        """One sentence naming L1 and the amount in figures and words."""
        l1 = next((v for v in self.data.vendors if v.rank == 1), None)
        if not l1 or l1.total is None:
            return
        para = self.document.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.add_run("L1 Bidder: ").bold = True
        para.add_run(
            f"{l1.name} at a total of {plain_amount(l1.total)}/- "
            f"({inr_words(l1.total)}) inclusive of GST @{self._gst_rate_text()}%."
        )

    def _item_wise_prices(self) -> None:
        """Per-line-item pricing for each ranked vendor.

        Only worth printing when the event has more than one priced line: with a
        single line this table would just restate the comparison above it.
        """
        items = [i for i in self.data.line_items if i.has_prices]
        vendors = [v for v in self.data.vendors if v.rank]
        if len(items) < 2 or not vendors:
            return

        self._heading("Item-wise Supplier Price Detail")

        headers = ["#", "Line Item", "Qty", "UoM"] + [
            f"{v.rank_label} — {v.name}" for v in vendors
        ]
        widths = ITEM_FIXED_WIDTHS + _vendor_column_widths(len(vendors))

        table = self.document.add_table(rows=len(items) + 2, cols=len(headers))
        table.style = "Table Grid"
        set_col_widths(table, widths)
        set_cell_margins(table)

        header_row = table.rows[0]
        repeat_header_row(header_row)
        for index, text in enumerate(headers):
            cell = header_row.cells[index]
            shade_cell(cell, styles.HEADER_FILL)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(8)

        totals = {v.name: None for v in vendors}
        for row_index, item in enumerate(items, start=1):
            row = table.rows[row_index]
            prevent_row_split(row)
            cells = [
                (str(row_index), WD_ALIGN_PARAGRAPH.CENTER),
                (item.title, WD_ALIGN_PARAGRAPH.LEFT),
                (plain_amount(item.quantity), WD_ALIGN_PARAGRAPH.RIGHT),
                (item.uom or PLACEHOLDER, WD_ALIGN_PARAGRAPH.CENTER),
            ]
            for vendor in vendors:
                amount = item.price_for(vendor.name)
                cells.append((plain_amount(amount), WD_ALIGN_PARAGRAPH.RIGHT))
                if amount is not None:
                    current = totals[vendor.name] or Decimal("0")
                    totals[vendor.name] = current + amount

            for index, (text, alignment) in enumerate(cells):
                para = row.cells[index].paragraphs[0]
                para.alignment = alignment
                run = para.add_run(text)
                run.font.size = Pt(8.5)

        self._item_totals_row(table.rows[-1], vendors, totals)
        self._item_wise_note(items, vendors)

    def _item_totals_row(self, row, vendors, totals) -> None:
        prevent_row_split(row)
        for index in range(4):
            shade_cell(row.cells[index], styles.HEADER_FILL)
        para = row.cells[1].paragraphs[0]
        para.add_run("Total").bold = True
        para.runs[0].font.size = Pt(8.5)

        for offset, vendor in enumerate(vendors, start=4):
            cell = row.cells[offset]
            shade_cell(cell, styles.HEADER_FILL)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run(plain_amount(totals.get(vendor.name)))
            run.bold = True
            run.font.size = Pt(8.5)

    def _item_wise_note(self, items, vendors) -> None:
        """Say plainly when a vendor did not quote every line."""
        gaps = [
            f"{v.name} ({sum(1 for i in items if i.price_for(v.name) is None)})"
            for v in vendors
            if any(i.price_for(v.name) is None for i in items)
        ]
        if not gaps:
            return
        para = self.document.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(
            "Lines not quoted (count in brackets): " + ", ".join(gaps)
            + ". Totals cover quoted lines only and are not directly comparable."
        )
        run.italic = True
        run.font.size = Pt(8.5)

    def _line_item_detail(self) -> None:
        """Item master detail, with the L1 price beside the last purchased price.

        Separate from the item-wise comparison above: this one answers "what is
        this line, and how does the winning price compare with what we last paid
        for it", which is the question an approver actually asks.
        """
        items = self.data.line_items
        if not items:
            return

        l1 = next((v for v in self.data.vendors if v.rank == 1), None)
        l1_name = l1.name if l1 else ""

        self._heading("Line Item Detail")

        headers = ["Item ID", "Item Name", "UOM", "Qty", "Last PO Price",
                   "L1 Price", "Total Price", "L1 Vendor"]
        table = self.document.add_table(rows=len(items) + 1, cols=len(headers))
        table.style = "Table Grid"
        set_col_widths(table, DETAIL_WIDTHS)
        set_cell_margins(table)

        header_row = table.rows[0]
        repeat_header_row(header_row)
        for index, text in enumerate(headers):
            cell = header_row.cells[index]
            shade_cell(cell, styles.HEADER_FILL)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            run.bold = True
            run.font.size = Pt(7)

        for offset, item in enumerate(items, start=1):
            row = table.rows[offset]
            prevent_row_split(row)
            l1_price = item.price_for(l1_name) if l1_name else None
            cells = [
                (item.material_code or item.item_id or PLACEHOLDER, WD_ALIGN_PARAGRAPH.LEFT),
                (item.title or PLACEHOLDER, WD_ALIGN_PARAGRAPH.LEFT),
                (item.uom or PLACEHOLDER, WD_ALIGN_PARAGRAPH.CENTER),
                (plain_amount(item.quantity), WD_ALIGN_PARAGRAPH.RIGHT),
                # Verbatim - Ariba stores this as text with its own currency.
                (item.last_po_price or PLACEHOLDER, WD_ALIGN_PARAGRAPH.RIGHT),
                (plain_amount(l1_price), WD_ALIGN_PARAGRAPH.RIGHT),
                (plain_amount(item.total_price), WD_ALIGN_PARAGRAPH.RIGHT),
                (l1_name if l1_price is not None else PLACEHOLDER, WD_ALIGN_PARAGRAPH.LEFT),
            ]
            for index, (text, alignment) in enumerate(cells):
                para = row.cells[index].paragraphs[0]
                para.alignment = alignment
                run = para.add_run(text)
                run.font.size = Pt(7.5)

        if not l1_name:
            para = self.document.add_paragraph()
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(
                "L1 price and vendor are blank because no priced offer has been "
                "received for this event yet."
            )
            run.italic = True
            run.font.size = Pt(8.5)

    def _justification(self) -> None:
        self._heading("Justification")
        text = (self.data.justification or "").strip()
        if not text:
            para = self.document.add_paragraph()
            self._missing_run(para, "Justification not provided")
            return
        for block in text.split("\n"):
            para = self.document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.add_run(block.strip())

    def _footer(self) -> None:
        footer = self.document.sections[0].footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_x_of_y(para)
        for run in para.runs:
            run.font.size = Pt(9)

    def _core_properties(self) -> None:
        props = self.document.core_properties
        props.title = f"NFA {self.data.pr_number}".strip()
        props.subject = self.data.subject or ""
        props.comments = f"Generated from Ariba event {self.data.event_id}".strip()

    # ------------------------------------------------------------------ #
    # Helpers
    # ------------------------------------------------------------------ #

    def _heading(self, text: str) -> None:
        para = self.document.add_paragraph()
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(text)
        run.bold = True
        run.underline = True

    def _missing_run(self, paragraph, text: str):
        run = paragraph.add_run(f"«{text}»")
        run.style = self.document.styles[styles.MISSING_STYLE]
        return run

    def _multiline(self, paragraph, text: str, size=None):
        """Render a value that may contain newlines as line breaks in one cell."""
        lines = (text or PLACEHOLDER).split("\n")
        for index, line in enumerate(lines):
            if index:
                paragraph.add_run().add_break()
            run = paragraph.add_run(line)
            if size:
                run.font.size = size

    def _gst_rate_text(self) -> str:
        rate = self.data.gst_rate
        return str(int(rate)) if rate == int(rate) else str(rate)


def build_nfa(data: NFAData) -> Document:
    return NFABuilder(data).build()


#: Subjects run to a couple of hundred characters; the reference file used a
#: short human descriptor. Trim to a word boundary so the name stays readable.
_SUBJECT_CHARS = 60


#: Trailing words that read as truncation damage rather than a short title.
_DANGLING = {
    "of", "for", "at", "in", "on", "to", "and", "or", "the", "a", "an",
    "with", "by", "from", "as", "-", "&",
}


def _short_subject(subject: str) -> str:
    text = re.sub(r"^NFA\s+(for\s+)?", "", subject.strip(), flags=re.IGNORECASE)
    text = re.sub(r"\s+", " ", text).rstrip(".")
    if len(text) <= _SUBJECT_CHARS:
        return text

    clipped = text[:_SUBJECT_CHARS]
    if " " in clipped:
        clipped = clipped[:clipped.rfind(" ")]
    words = clipped.rstrip(" ,;-").split(" ")
    while len(words) > 1 and words[-1].lower().strip(",;-") in _DANGLING:
        words.pop()
    return " ".join(words).rstrip(" ,;-")


def output_filename(data: NFAData) -> str:
    """Match the existing naming convention: 'NFA <PR number> <subject>.docx'."""
    parts = ["NFA"]
    if data.pr_number:
        parts.append(str(data.pr_number))
    elif data.event_id:
        parts.append(str(data.event_id))
    if data.subject:
        parts.append(_short_subject(data.subject))
    name = " ".join(p for p in parts if p)
    name = re.sub(r'[<>:"/\\|?*]', "", name)      # characters Windows forbids
    name = re.sub(r"\s+", " ", name).strip()
    return f"{name}.docx"
