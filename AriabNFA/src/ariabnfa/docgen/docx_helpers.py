"""Low-level OOXML helpers for things python-docx has no API for.

Each function here exists because the reference NFA needs a feature the library
does not expose: the double page border, PAGE/NUMPAGES fields in the footer,
genuinely fixed column widths, and repeating table header rows.

OOXML child elements are order-sensitive. Word will often repair a document
with elements in the wrong order, but not always, and a "repaired" prompt on an
approval document looks like corruption to the person opening it. So insertions
go through `insert_in_order` with the schema sequence spelled out.
"""

from __future__ import annotations

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

# Schema sequences, truncated to the elements that follow our insertion point.
_SECTPR_AFTER_PGBORDERS = [
    "lnNumType", "pgNumType", "cols", "formProt", "vAlign", "noEndnote",
    "titlePg", "textDirection", "bidi", "rtlGutter", "docGrid",
    "printerSettings", "sectPrChange",
]
_TBLPR_AFTER_TBLLAYOUT = ["tblCellMar", "tblLook", "tblCaption", "tblDescription", "tblPrChange"]
_TCPR_AFTER_SHD = [
    "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign", "hideMark",
]


def insert_in_order(parent, element, following_tags: list[str]) -> None:
    """Insert `element` before the first of `following_tags` present in `parent`.

    Falls back to appending when none of them are there.
    """
    for tag in following_tags:
        existing = parent.find(qn(f"w:{tag}"))
        if existing is not None:
            existing.addprevious(element)
            return
    parent.append(element)


def _replace_child(parent, tag: str):
    existing = parent.find(qn(tag))
    if existing is not None:
        parent.remove(existing)
    return OxmlElement(tag)


# --------------------------------------------------------------------------- #
# Page setup
# --------------------------------------------------------------------------- #

def set_page_borders(
    section,
    *,
    style: str = "double",
    size: int = 4,
    space: int = 20,
    color: str = "auto",
) -> None:
    """Draw a border around the whole page on all four sides.

    The reference NFA uses a double line offset from the page edge; without it
    the document reads as visibly not-the-usual-form at a glance.
    """
    sect_pr = section._sectPr
    borders = _replace_child(sect_pr, "w:pgBorders")
    borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), style)
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), str(space))
        el.set(qn("w:color"), color)
        borders.append(el)
    insert_in_order(sect_pr, borders, _SECTPR_AFTER_PGBORDERS)


# --------------------------------------------------------------------------- #
# Fields
# --------------------------------------------------------------------------- #

def add_field(paragraph, instruction: str, placeholder: str = "1"):
    """Append a Word field such as PAGE or NUMPAGES to `paragraph`.

    Word computes the real value when it repaginates, so a freshly generated
    file may show the placeholder until then. That is expected and resolves on
    print preview or any edit.
    """
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = placeholder

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for el in (begin, instr, separate, text, end):
        run._r.append(el)
    return run


def add_page_x_of_y(paragraph, *, prefix: str = "Page ", joiner: str = " of ") -> None:
    """Render 'Page X of Y' using real PAGE and NUMPAGES fields."""
    paragraph.add_run(prefix)
    add_field(paragraph, "PAGE", "1")
    paragraph.add_run(joiner)
    add_field(paragraph, "NUMPAGES", "1")


# --------------------------------------------------------------------------- #
# Tables
# --------------------------------------------------------------------------- #

def set_fixed_layout(table) -> None:
    """Switch a table to fixed layout so declared widths are actually honoured."""
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = _replace_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    insert_in_order(tbl_pr, layout, _TBLPR_AFTER_TBLLAYOUT)


def set_col_widths(table, widths_cm: list[float]) -> None:
    """Fix column widths.

    Setting `table.columns[i].width` alone does not work - Word takes the width
    from the individual cells. The width has to be written to every cell in the
    column *and* the table put into fixed layout, or the columns silently
    resize to fit their content.
    """
    set_fixed_layout(table)
    for index, width in enumerate(widths_cm):
        if index >= len(table.columns):
            break
        table.columns[index].width = Cm(width)
        for cell in table.columns[index].cells:
            cell.width = Cm(width)


def repeat_header_row(row) -> None:
    """Mark a row as a header that repeats on every page the table spans."""
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:tblHeader"))
    if existing is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def shade_cell(cell, fill: str) -> None:
    """Apply a solid background fill to a cell, e.g. 'D9D9D9'."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = _replace_child(tc_pr, "w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    insert_in_order(tc_pr, shd, _TCPR_AFTER_SHD)


def set_cell_margins(table, *, top=0.08, bottom=0.08, left=0.15, right=0.15) -> None:
    """Set default cell padding (in cm) for every cell in the table."""
    tbl_pr = table._tbl.tblPr
    margins = _replace_child(tbl_pr, "w:tblCellMar")
    for tag, value in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{tag}")
        el.set(qn("w:w"), str(int(value * 567)))  # cm -> twips
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    insert_in_order(tbl_pr, margins, ["tblLook", "tblCaption", "tblDescription", "tblPrChange"])


def remove_table_borders(table) -> None:
    """Strip all borders - used for the letterhead block, which is a layout
    table rather than a data table."""
    tbl_pr = table._tbl.tblPr
    borders = _replace_child(tbl_pr, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        borders.append(el)
    insert_in_order(tbl_pr, borders, ["shd", "tblLayout", "tblCellMar", "tblLook"])


def prevent_row_split(row) -> None:
    """Keep a table row from breaking across a page."""
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))
